"""
Deep analysis of docx formatting — paragraph by paragraph, run by run.
Outputs a comprehensive format specification as JSON and a human-readable summary.

Usage:
    python analyze_docx_format.py --input <docx_path> [--output-dir <dir>]
"""
from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import argparse
import json
import os


def emu_to_pt(emu):
    """Convert EMU to points."""
    if emu is None:
        return None
    return round(emu / 12700, 1)


def emu_to_cm(emu):
    """Convert EMU to centimeters."""
    if emu is None:
        return None
    return round(emu / 360000, 2)


def rgb_to_hex(rgb):
    """Convert RGBColor to hex string."""
    if rgb is None:
        return None
    try:
        return str(rgb)
    except Exception:
        return None


def analyze_paragraph(para, idx):
    """Deep analysis of a single paragraph."""
    info = {
        'index': idx,
        'style': para.style.name if para.style else None,
        'text_preview': para.text[:100] if para.text else '',
        'paragraph_format': {
            'alignment': str(para.alignment) if para.alignment else None,
            'line_spacing': para.paragraph_format.line_spacing,
            'line_spacing_rule': str(para.paragraph_format.line_spacing_rule)
                if para.paragraph_format.line_spacing_rule else None,
            'space_before': emu_to_pt(para.paragraph_format.space_before),
            'space_after': emu_to_pt(para.paragraph_format.space_after),
            'first_line_indent': emu_to_cm(para.paragraph_format.first_line_indent),
            'left_indent': emu_to_cm(para.paragraph_format.left_indent),
            'right_indent': emu_to_cm(para.paragraph_format.right_indent),
        },
        'runs': [
            {
                'text_preview': run.text[:80] if run.text else '',
                'font_name': run.font.name,
                'font_size_pt': emu_to_pt(run.font.size),
                'bold': run.font.bold,
                'italic': run.font.italic,
                'underline': run.font.underline,
                'color_hex': rgb_to_hex(run.font.color.rgb)
                    if run.font.color and run.font.color.rgb else None,
            }
            for run in para.runs
        ],
    }
    return info


def analyze_table(table, idx):
    """Deep analysis of a table."""
    info = {
        'index': idx,
        'rows': len(table.rows),
        'columns': len(table.columns),
    }

    rows_info = []
    for ri, row in enumerate(table.rows):
        cells_info = []
        for ci, cell in enumerate(row.cells):
            cinfo = {
                'text': cell.text[:100].replace('\n', ' '),
                'paragraphs': [
                    analyze_paragraph(para, pi)
                    for pi, para in enumerate(cell.paragraphs)
                    if para.text.strip()
                ],
            }
            cells_info.append(cinfo)
        rows_info.append(cells_info)
        if ri > 20:  # Limit output
            break

    info['rows_data'] = rows_info
    return info


def analyze_docx(doc_path, output_dir):
    """Analyze a docx file and write format spec to output_dir."""
    doc = Document(doc_path)

    # Analyze sections
    sections_info = []
    for si, section in enumerate(doc.sections):
        sinfo = {
            'index': si,
            'page_width_cm': emu_to_cm(section.page_width),
            'page_height_cm': emu_to_cm(section.page_height),
            'margin_top_cm': emu_to_cm(section.top_margin),
            'margin_bottom_cm': emu_to_cm(section.bottom_margin),
            'margin_left_cm': emu_to_cm(section.left_margin),
            'margin_right_cm': emu_to_cm(section.right_margin),
            'orientation': str(section.orientation) if section.orientation else None,
        }
        if section.header:
            sinfo['header_text'] = ' '.join(
                p.text for p in section.header.paragraphs
            )
            sinfo['header_paragraphs'] = [
                analyze_paragraph(p, i)
                for i, p in enumerate(section.header.paragraphs)
                if p.text.strip()
            ]
        if section.footer:
            sinfo['footer_text'] = ' '.join(
                p.text for p in section.footer.paragraphs
            )
            sinfo['footer_paragraphs'] = [
                analyze_paragraph(p, i)
                for i, p in enumerate(section.footer.paragraphs)
                if p.text.strip()
            ]
        sections_info.append(sinfo)

    # Analyze all paragraphs
    paragraphs_info = [
        analyze_paragraph(para, i)
        for i, para in enumerate(doc.paragraphs)
        if para.text.strip()
    ]

    # Analyze tables
    tables_info = [
        analyze_table(table, ti)
        for ti, table in enumerate(doc.tables)
    ]

    # Build format spec
    format_spec = {
        'source_file': os.path.basename(doc_path),
        'sections': sections_info,
        'paragraph_count': len(paragraphs_info),
        'table_count': len(tables_info),
    }

    # Write JSON
    json_path = os.path.join(output_dir, 'format_analysis.json')
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(format_spec, f, ensure_ascii=False, indent=2, default=str)

    # Write human-readable summary
    summary_path = os.path.join(output_dir, 'format_summary.txt')
    with open(summary_path, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write(f"标书格式规范总结: {os.path.basename(doc_path)}\n")
        f.write("=" * 80 + "\n\n")

        f.write("【页面设置】\n")
        for s in sections_info:
            f.write(f"  页面尺寸: {s['page_width_cm']}cm × {s['page_height_cm']}cm (A4)\n")
            f.write(f"  页边距: 上={s['margin_top_cm']}cm 下={s['margin_bottom_cm']}cm "
                    f"左={s['margin_left_cm']}cm 右={s['margin_right_cm']}cm\n")
            if s.get('header_text'):
                f.write(f"  页眉: {s['header_text']}\n")
            if s.get('footer_text'):
                f.write(f"  页脚: {s['footer_text']}\n")

        f.write(f"\n【正文段落数】{len(paragraphs_info)}\n")
        f.write(f"【表格数】{len(tables_info)}\n")

        f.write("\n【内容结构模式】\n")
        f.write("  每个服务响应采用统一模式:\n")
        f.write("  【招标要求】+ 招标原文引用\n")
        f.write("  【我方响应】+ 详细实施方案\n")
        f.write("  响应内容: 具体、量化、可操作，避免空泛承诺\n")

    print(f'Format analysis written to:\n  {json_path}\n  {summary_path}')


def main():
    parser = argparse.ArgumentParser(
        description='Deep analysis of docx formatting for bid documents'
    )
    parser.add_argument('--input', '-i', required=True,
                        help='Path to the input .docx file to analyze')
    parser.add_argument('--output-dir', '-o', default=None,
                        help='Output directory for format_analysis.json and format_summary.txt '
                             '(default: same directory as the script)')
    args = parser.parse_args()

    output_dir = args.output_dir or os.path.dirname(os.path.abspath(__file__))
    os.makedirs(output_dir, exist_ok=True)

    analyze_docx(args.input, output_dir)


if __name__ == '__main__':
    main()

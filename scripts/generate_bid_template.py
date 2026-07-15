"""
Generate a bid document template (.docx)

This template preserves all formatting: fonts, sizes, colors, alignment,
spacing, headings, tables, page layout, headers/footers.

Usage:
    python generate_bid_template.py --output <output_path.docx>
"""
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import argparse
from datetime import datetime


def set_run_font(run, cjk_font, latin_font='Times New Roman', font_size=None, bold=None, color=None):
    """
    Set font properties on a run with separate Latin and CJK fonts.

    OOXML uses three font slots:
    - w:ascii / w:hAnsi → Latin/Western characters (Times New Roman)
    - w:eastAsia        → CJK characters (宋体/黑体)

    python-docx's run.font.name only sets w:ascii and w:hAnsi, so we
    set that to the Latin font, then patch w:eastAsia via XML for CJK.
    """
    run.font.name = latin_font
    if font_size is not None:
        run.font.size = font_size
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color

    # Patch the underlying XML to add w:eastAsia for CJK
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cjk_font)
    # Also set hAnsi explicitly (run.font.name sets ascii, we ensure hAnsi too)
    rFonts.set(qn('w:hAnsi'), latin_font)


def set_cell_font(cell, text, cjk_font='宋体', latin_font='Times New Roman',
                  font_size=None, bold=False, alignment=None, color=None):
    """Set font properties for a table cell, with separate Latin/CJK fonts."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    para = cell.paragraphs[0]
    if alignment:
        para.alignment = alignment
    run = para.add_run(str(text))
    set_run_font(run, cjk_font, latin_font=latin_font, font_size=font_size, bold=bold, color=color)
    return run


class BidTemplateGenerator:
    """Generate a bid document template matching the sample format."""

    # === Format Constants (derived from format analysis) ===
    # Page
    PAGE_WIDTH = Cm(21.0)   # A4
    PAGE_HEIGHT = Cm(29.7)  # A4
    MARGIN_TOP = Cm(2.5)
    MARGIN_BOTTOM = Cm(2.5)
    MARGIN_LEFT = Cm(3.0)
    MARGIN_RIGHT = Cm(2.5)

    # Fonts — Chinese fonts (CJK) and Latin font for Western characters
    FONT_LATIN = 'Times New Roman'     # All Western/Latin characters
    FONT_BODY = '宋体'                  # Body text CJK
    FONT_HEADING = '黑体'               # All headings and cover title CJK
    FONT_COVER_TITLE = '黑体'           # Cover main title CJK (= FONT_HEADING)
    FONT_COVER_INFO = '宋体'            # Cover info lines CJK (= FONT_BODY)
    FONT_TABLE_HEADER = '黑体'          # Table header CJK (= FONT_HEADING)

    # Colors
    COLOR_HEADING1 = RGBColor(0x37, 0x60, 0x92)  # #376092
    COLOR_HEADING2 = RGBColor(0x4F, 0x81, 0xBD)  # #4F81BD
    COLOR_HEADING3 = RGBColor(0x4F, 0x81, 0xBD)  # #4F81BD
    COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)

    # Sizes (in Pt)
    SIZE_COVER_TITLE = Pt(36)
    SIZE_COVER_SUBTITLE = Pt(16)
    SIZE_COVER_INFO = Pt(14)
    SIZE_BODY = Pt(12)
    SIZE_H1 = Pt(16)
    SIZE_H2 = Pt(15)
    SIZE_H3 = Pt(14)
    SIZE_TABLE = Pt(10.5)

    # Spacing
    LINE_SPACING = 1.5
    BODY_SPACE_BEFORE = Pt(6)
    BODY_SPACE_AFTER = Pt(6)
    H1_SPACE_BEFORE = Pt(12)
    H1_SPACE_AFTER = Pt(6)
    H2_SPACE_BEFORE = Pt(10)
    H2_SPACE_AFTER = Pt(6)
    H3_SPACE_BEFORE = Pt(8)
    H3_SPACE_AFTER = Pt(4)

    def __init__(self):
        self.doc = Document()

    def setup_page(self):
        """Configure page size, margins, and orientation."""
        for section in self.doc.sections:
            section.page_width = self.PAGE_WIDTH
            section.page_height = self.PAGE_HEIGHT
            section.top_margin = self.MARGIN_TOP
            section.bottom_margin = self.MARGIN_BOTTOM
            section.left_margin = self.MARGIN_LEFT
            section.right_margin = self.MARGIN_RIGHT

    def setup_styles(self):
        """Configure document styles: CJK fonts (宋体/黑体), Latin font (Times New Roman)."""
        # Normal style: CJK=宋体, Latin=Times New Roman
        style = self.doc.styles['Normal']
        style.font.name = self.FONT_LATIN  # Latin font
        style.font.size = self.SIZE_COVER_INFO
        style.font.color.rgb = self.COLOR_BLACK
        style.paragraph_format.line_spacing = self.LINE_SPACING
        self._patch_style_eastAsia('Normal', self.FONT_BODY)  # CJK=宋体

        # Heading 1: CJK=黑体, Latin=Times New Roman
        h1 = self.doc.styles['Heading 1']
        h1.font.name = self.FONT_LATIN
        h1.font.size = self.SIZE_H1
        h1.font.bold = True
        h1.font.color.rgb = self.COLOR_HEADING1
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1.paragraph_format.line_spacing = self.LINE_SPACING
        h1.paragraph_format.space_before = self.H1_SPACE_BEFORE
        h1.paragraph_format.space_after = self.H1_SPACE_AFTER
        self._patch_style_eastAsia('Heading 1', self.FONT_HEADING)

        # Heading 2: CJK=黑体, Latin=Times New Roman
        h2 = self.doc.styles['Heading 2']
        h2.font.name = self.FONT_LATIN
        h2.font.size = self.SIZE_H2
        h2.font.bold = True
        h2.font.color.rgb = self.COLOR_HEADING2
        h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2.paragraph_format.line_spacing = self.LINE_SPACING
        h2.paragraph_format.space_before = self.H2_SPACE_BEFORE
        h2.paragraph_format.space_after = self.H2_SPACE_AFTER
        self._patch_style_eastAsia('Heading 2', self.FONT_HEADING)

        # Heading 3: CJK=黑体, Latin=Times New Roman
        h3 = self.doc.styles['Heading 3']
        h3.font.name = self.FONT_LATIN
        h3.font.size = self.SIZE_H3
        h3.font.bold = True
        h3.font.color.rgb = self.COLOR_HEADING3
        h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h3.paragraph_format.line_spacing = self.LINE_SPACING
        h3.paragraph_format.space_before = self.H3_SPACE_BEFORE
        h3.paragraph_format.space_after = self.H3_SPACE_AFTER
        self._patch_style_eastAsia('Heading 3', self.FONT_HEADING)

        # Patch theme eastAsia fonts (also empty by default)
        self._patch_theme_eastAsia()

    def _patch_style_eastAsia(self, style_name, font_name):
        """Patch a style's rPr to include w:eastAsia font attribute."""
        style = self.doc.styles[style_name]
        rPr = style.element.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            style.element.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)

    def _patch_theme_eastAsia(self):
        """Patch document theme to include eastAsia font names."""
        # Access the theme part via the document part
        try:
            theme_part = self.doc.part.package.part_related_by(
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme'
            )
            if hasattr(theme_part, '_element'):
                theme_el = theme_part._element
            else:
                import io
                theme_el = OxmlElement('placeholder')
                return
        except Exception:
            return

        ns_a = 'http://schemas.openxmlformats.org/drawingml/2006/main'
        for font_scheme in theme_el.iter(f'{{{ns_a}}}fontScheme'):
            # Patch majorFont
            major = font_scheme.find(f'{{{ns_a}}}majorFont')
            if major is not None:
                ea = major.find(f'{{{ns_a}}}ea')
                if ea is None:
                    ea = OxmlElement(f'{{{ns_a}}}ea')
                    major.insert(0, ea)
                ea.set('typeface', self.FONT_HEADING)  # 黑体 for headings
            # Patch minorFont
            minor = font_scheme.find(f'{{{ns_a}}}minorFont')
            if minor is not None:
                ea = minor.find(f'{{{ns_a}}}ea')
                if ea is None:
                    ea = OxmlElement(f'{{{ns_a}}}ea')
                    minor.insert(0, ea)
                ea.set('typeface', self.FONT_BODY)  # 宋体 for body

    def add_cover_paragraph(self, text, cjk_font=None, font_size=None, bold=False, alignment=None):
        """Add a paragraph with cover page formatting.
        cjk_font: CJK font (黑体 for title, 宋体 for info lines). Defaults to 宋体."""
        para = self.doc.add_paragraph()
        if alignment:
            para.alignment = alignment
        run = para.add_run(text)
        fn = cjk_font or self.FONT_COVER_INFO  # Default to 宋体 for info lines
        set_run_font(run, fn, latin_font=self.FONT_LATIN,
                     font_size=font_size or self.SIZE_COVER_INFO,
                     bold=bold, color=self.COLOR_BLACK)
        para.paragraph_format.line_spacing = self.LINE_SPACING
        return para

    def add_body_paragraph(self, text, bold=False):
        """Add a body text paragraph with standard formatting.
        CJK: 宋体, Latin: Times New Roman."""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(text)
        set_run_font(run, self.FONT_BODY, latin_font=self.FONT_LATIN,
                     font_size=self.SIZE_BODY, bold=bold, color=self.COLOR_BLACK)
        para.paragraph_format.line_spacing = self.LINE_SPACING
        para.paragraph_format.space_before = self.BODY_SPACE_BEFORE
        para.paragraph_format.space_after = self.BODY_SPACE_AFTER
        return para

    def _add_styled_heading(self, text, style_name, cjk_font, font_size, bold, color):
        """Add a heading paragraph with explicit run-level fonts.

        Must set w:eastAsia directly on the run (not just the style) because
        WPS/Word may skip style-level CJK font definitions and fall back to
        the theme's majorFont (defaulting to MS Gothic if unset).
        """
        para = self.doc.add_paragraph(style=style_name)
        # Clear the default run that add_paragraph creates when text is passed
        run = para.add_run(text)
        set_run_font(run, cjk_font, latin_font=self.FONT_LATIN,
                     font_size=font_size, bold=bold, color=color)
        return para

    def add_heading1(self, text):
        """Add a Heading 1 paragraph."""
        return self._add_styled_heading(
            text, 'Heading 1', self.FONT_HEADING, self.SIZE_H1,
            bold=True, color=self.COLOR_HEADING1)

    def add_heading2(self, text):
        """Add a Heading 2 paragraph."""
        return self._add_styled_heading(
            text, 'Heading 2', self.FONT_HEADING, self.SIZE_H2,
            bold=True, color=self.COLOR_HEADING2)

    def add_heading3(self, text):
        """Add a Heading 3 paragraph."""
        return self._add_styled_heading(
            text, 'Heading 3', self.FONT_HEADING, self.SIZE_H3,
            bold=True, color=self.COLOR_HEADING3)

    def add_requirement_response(self, requirement, response):
        """Add a 【招标要求】... 【我方响应】... pattern."""
        self.add_body_paragraph(f'【招标要求】{requirement}')
        self.add_body_paragraph(f'【我方响应】{response}')

    def add_table_from_data(self, headers, rows, col_widths=None):
        """Add a formatted table with header row and data rows."""
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'

        # Header row
        for ci, header in enumerate(headers):
            cell = table.rows[0].cells[ci]
            set_cell_font(cell, header,
                          cjk_font=self.FONT_TABLE_HEADER,
                          latin_font=self.FONT_LATIN,
                          font_size=self.SIZE_TABLE,
                          bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
            # Light blue background
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D5E8F0')
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)

        # Data rows
        for ri, row_data in enumerate(rows):
            for ci, cell_text in enumerate(row_data):
                cell = table.rows[ri + 1].cells[ci]
                set_cell_font(cell, str(cell_text) if cell_text else '',
                              cjk_font=self.FONT_BODY,
                              latin_font=self.FONT_LATIN,
                              font_size=self.SIZE_TABLE)

        return table

    def add_page_break(self):
        """Add a page break."""
        self.doc.add_page_break()

    # === Template Generation ===

    def generate_cover_page(self):
        """Generate the cover (first) page.
        Title: 黑体 + Times New Roman
        Info lines: 宋体 + Times New Roman"""
        # Empty lines for spacing
        for _ in range(3):
            self.add_cover_paragraph('')

        # Main title — 黑体 36pt
        self.add_cover_paragraph('投  标  文  件',
                                 cjk_font=self.FONT_COVER_TITLE,
                                 font_size=self.SIZE_COVER_TITLE,
                                 bold=True,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER)

        self.add_cover_paragraph('')

        # Project name — 黑体 16pt
        self.add_cover_paragraph('项目名称：【项目名称】',
                                 cjk_font=self.FONT_COVER_TITLE,
                                 font_size=self.SIZE_COVER_SUBTITLE,
                                 bold=True,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER)

        self.add_cover_paragraph('')

        # Bidder info lines — 宋体 14pt
        self.add_cover_paragraph('投标人：（公章）【投标人全称】',
                                 cjk_font=self.FONT_COVER_INFO,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self.add_cover_paragraph('法定代表人或其委托代理人：（签字）',
                                 cjk_font=self.FONT_COVER_INFO,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self.add_cover_paragraph('地址：【投标人地址】',
                                 cjk_font=self.FONT_COVER_INFO,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self.add_cover_paragraph('联系电话：【联系电话】',
                                 cjk_font=self.FONT_COVER_INFO,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self.add_cover_paragraph('日  期：【    年    月    日】',
                                 cjk_font=self.FONT_COVER_INFO,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER)

    def generate_toc(self):
        """Generate the table of contents page."""
        self.add_page_break()
        self.add_heading1('目  录')
        self.add_body_paragraph('（此处由Word自动生成目录，请在Word中右键点击此处选择"更新域"）')

    def generate_section1_proposal(self):
        """Generate Section 1: 采购方案 (Technical Proposal)."""
        self.add_page_break()
        self.add_heading1('一、采购方案')

        # 1.1 项目理解
        self.add_heading2('1.1 项目理解')
        self.add_body_paragraph(
            '【描述对项目的整体理解，包括：项目背景、项目名称、采购单位、预算金额、服务周期、'
            '核心需求分析、评审方式分析等。参考模式：本项目为"【项目名称】"，由【采购单位】'
            '组织自行采购，预算金额【金额】万元，服务周期为【周期】。项目核心需求分为【N】大板块：'
            '一是【需求一】；二是【需求二】；三是【需求三】。项目评审采用综合评分法，'
            '【分析各评分项权重】。】'
        )

        # 1.2 供应商基本情况表
        self.add_heading2('1.2 供应商基本情况表')
        self.add_body_paragraph('按招标文件附件1要求，供应商基本情况表如下：')
        self.add_table_from_data(
            ['项目', '内容'],
            [
                ['供应商名称', '【投标人全称】'],
                ['统一社会信用代码', '【统一社会信用代码】'],
                ['注册地址', '【注册地址】'],
                ['法定代表人', '【法定代表人姓名】'],
                ['注册资本', '【注册资本】'],
                ['成立日期', '【成立日期】'],
                ['经营范围', '【经营范围】'],
                ['联系人', '【联系人姓名】'],
                ['联系电话', '【联系电话】'],
                ['电子邮箱', '【电子邮箱】'],
                ['开户银行', '【开户银行】'],
                ['银行账号', '【银行账号】'],
            ]
        )

        # 1.3+ Service response sections (template placeholders)
        self.add_heading2('1.3 【服务模块一名称】')
        self.add_heading3('1.3.1 【子服务名称】')
        self.add_body_paragraph('针对招标文件相关要求，我方逐条响应如下：')
        self.add_requirement_response(
            '【此处填写招标要求原文】',
            '【此处填写我方详细响应方案，必须：1.逐条对应招标要求；2.具体、量化、可操作；'
            '3.说明实施方法、流程、交付物、时间节点；4.避免空泛承诺】'
        )

        self.add_heading2('1.4 【服务模块二名称】')
        self.add_heading3('1.4.1 【子服务名称】')
        self.add_requirement_response(
            '【此处填写招标要求原文】',
            '【此处填写我方详细响应方案】'
        )

        # Work method
        self.add_heading2('1.6 履行采购项目的工作方式')
        self.add_body_paragraph(
            '（1）驻场服务模式。【描述驻场人员配置、工作地点、工作时间等】\n'
            '（2）远程服务支撑。【描述远程支撑体系】\n'
            '（3）定期沟通机制。【描述日/周/月沟通机制】\n'
            '（4）工单驱动。【描述工单管理系统】'
        )

        # Fee breakdown
        self.add_heading2('1.7 费用明细')
        self.add_body_paragraph(
            '本项目总报价：【待补充】万元（大写：【待补充】元整）。'
            '预算金额【金额】万元，总报价不超过预算。'
        )
        self.add_table_from_data(
            ['序号', '费用项目', '服务内容', '报价（万元）', '备注'],
            [
                ['1', '【费用项目一】', '【服务内容描述】', '【待补充】', '详见对应章节'],
                ['2', '【费用项目二】', '【服务内容描述】', '【待补充】', '详见对应章节'],
                ['3', '【费用项目三】', '【服务内容描述】', '【待补充】', '详见对应章节'],
                ['', '合  计', '', '【待补充】', ''],
            ]
        )

    def generate_section2_qualifications(self):
        """Generate Section 2: 供应商资质证明材料及联系方式."""
        self.add_page_break()
        self.add_heading1('二、供应商资质证明材料及联系方式')
        self.add_body_paragraph('（1）营业执照（副本）复印件：见附件。')
        self.add_body_paragraph('（2）资质证书清单：')
        self.add_table_from_data(
            ['序号', '证书名称', '颁发机构', '证书编号', '有效期'],
            [
                ['1', '营业执照', '【颁发机构】', '【统一社会信用代码】', '【有效期】'],
                ['2', '【资质证书一】', '【颁发机构】', '【证书编号】', '【有效期】'],
                ['3', '【资质证书二】', '【颁发机构】', '【证书编号】', '【有效期】'],
            ]
        )
        self.add_body_paragraph('（3）联系方式：')
        self.add_table_from_data(
            ['项目', '内容'],
            [
                ['供应商名称', '【投标人全称】'],
                ['地址', '【投标人地址】'],
                ['联系人', '【联系人姓名】'],
                ['联系电话', '【联系电话】'],
                ['电子邮箱', '【电子邮箱】'],
                ['传真', '【传真号码，如无可填"无"】'],
            ]
        )
        self.add_body_paragraph('注：以上资质证书复印件请附后，原件备查。')

    def generate_section3_legal_rep_id(self):
        """Generate Section 3: 法定代表人身份证复印件."""
        self.add_page_break()
        self.add_heading1('三、公司法定代表人的身份证复印件')
        self.add_body_paragraph('公司法定代表人：')
        self.add_body_paragraph('姓名：【法定代表人姓名】')
        self.add_body_paragraph('身份证号码：【身份证号码】')
        self.add_body_paragraph('（此处粘贴法定代表人身份证正反面复印件，加盖公章）')

    def generate_section4_authorization(self):
        """Generate Section 4: 法定代表人授权委托书."""
        self.add_page_break()
        self.add_heading1('四、法定代表人授权委托书')
        self.add_body_paragraph('法定代表人授权委托书', bold=True)
        self.add_body_paragraph('致：【采购单位名称】')
        self.add_body_paragraph(
            '我，【法定代表人姓名】，系【投标人全称】的法定代表人，现授权委托'
            '【被授权人姓名】（身份证号码：【被授权人身份证号码】，职务：【被授权人职务】）'
            '为我方代理人，以我方名义参加"【项目名称】"的投标活动。代理人在投标、开标、'
            '评标、合同签署、合同执行过程中所签署的一切文件和处理与之有关的一切事务，'
            '我方均予以承认。'
        )
        self.add_body_paragraph('代理人无转委托权。')
        self.add_body_paragraph('本授权书自签署之日起生效，至本项目合同履行完毕之日止失效。')
        self.add_body_paragraph('')
        self.add_body_paragraph('投标人（公章）：【投标人全称】')
        self.add_body_paragraph('法定代表人（签字）：_______________')
        self.add_body_paragraph('被授权人（签字）：_______________')
        self.add_body_paragraph('日期：【    年    月    日】')
        self.add_body_paragraph('附：')
        self.add_body_paragraph('1. 法定代表人身份证复印件（加盖公章）')
        self.add_body_paragraph('2. 被授权人身份证复印件（加盖公章）')

    def generate_section5_commitment(self):
        """Generate Section 5: 符合资质要求的承诺函."""
        self.add_page_break()
        self.add_heading1('五、符合资质要求的承诺函')
        self.add_body_paragraph('承  诺  函', bold=True)
        self.add_body_paragraph('致：【采购单位名称】')
        self.add_body_paragraph('我单位（【投标人全称】）郑重声明并承诺如下：')
        self.add_body_paragraph('一、符合《中华人民共和国政府采购法》第二十二条规定的资质要求：')
        items = [
            '具有独立承担民事责任的能力；',
            '具有良好的商业信誉和健全的财务会计制度；',
            '具有履行合同所必需的设备和专业技术能力；',
            '有依法缴纳税收和社会保障资金的良好记录；',
            '参加政府采购活动前三年内，在经营活动中没有重大违法记录；',
            '法律、行政法规规定的其他条件。',
        ]
        for item in items:
            self.add_body_paragraph(f'    （{items.index(item)+1}）{item}')
        self.add_body_paragraph('二、参与本项目政府采购活动时不存在被有关部门禁止参与政府采购活动且在有效期内的情况。')
        self.add_body_paragraph('三、未被列入失信被执行人、重大税收违法案件当事人名单、政府采购严重违法失信行为记录名单。')
        self.add_body_paragraph('四、不存在《深圳市财政局政府采购供应商信用信息管理办法》（深财规〔2023〕3号）列明的严重违法失信行为。')
        self.add_body_paragraph('五、单位负责人为同一人或者存在直接控股、管理关系的不同供应商，不得参加同一合同项下的政府采购活动；我单位未为本采购项目提供整体设计、规范编制或者项目管理、监理、检测等服务。')
        self.add_body_paragraph('我单位对上述承诺声明的真实性负责。若有虚假承诺，视同提供虚假资料，自愿依法承担相应法律责任。')
        self.add_body_paragraph('')
        self.add_body_paragraph('承诺人（公章）：【投标人全称】')
        self.add_body_paragraph('法定代表人（签字）：_______________')
        self.add_body_paragraph('日期：【    年    月    日】')

    def generate_section6_service_evidence(self):
        """Generate Section 6: 符合项目服务要求的证明材料."""
        self.add_page_break()
        self.add_heading1('六、符合项目服务要求的证明材料')

        self.add_heading2('6.1 项目团队配置')
        self.add_body_paragraph('根据评分标准"项目团队综合能力"的要求，我方为本项目配备以下团队：')

        self.add_heading3('6.1.1 项目负责人（1人）')
        self.add_table_from_data(
            ['项目', '内容'],
            [
                ['姓名', '【姓名】'],
                ['学历', '计算机相关专业本科及以上学历'],
                ['PMP认证', '【证书编号】'],
                ['ITIL认证', '【证书编号】'],
                ['信息系统项目管理师（高级）', '【证书编号，由人力资源和社会保障部门颁发】'],
                ['社保', '见附件（开标日前由社保部门盖章的缴纳证明）'],
            ]
        )

        self.add_heading3('6.1.2 项目团队成员（3人）')
        self.add_table_from_data(
            ['姓名', '角色', '资质证书', '证书编号', '社保'],
            [
                ['【姓名】', '系统分析师/信息系统项目管理师', '高级证书', '【证书编号】', '见附件'],
                ['【姓名】', '企业微信运维工程师', '企业微信私有化运维工程师（中级）', '【证书编号】', '见附件'],
                ['【姓名】', 'ITIL运维管理', 'ITIL认证证书', '【证书编号】', '见附件'],
            ]
        )
        self.add_body_paragraph('注：以上人员相关有效证书复印件及社保缴纳证明材料请附后，原件备查。')

        self.add_heading2('6.2 同类项目业绩')
        self.add_body_paragraph('根据评分标准"同类项目业绩"的要求，投标人提供自【开始日期】起至本项目投标截止日前所承担的同类项目业绩情况。')
        self.add_table_from_data(
            ['序号', '项目名称', '业主单位', '合同金额(万元)', '合同签订时间', '服务内容'],
            [
                ['1', '【项目名称】', '【业主单位】', '【金额】', '【日期】', '【相关服务内容】'],
                ['2', '【项目名称】', '【业主单位】', '【金额】', '【日期】', '【相关服务内容】'],
                ['3', '【项目名称】', '【业主单位】', '【金额】', '【日期】', '【相关服务内容】'],
            ]
        )
        self.add_body_paragraph('注：以上项目合同关键页复印件请附后，原件备查。')

        self.add_heading2('6.3 公司资质证书')
        self.add_body_paragraph('根据评分标准"公司实力"的要求：')
        self.add_table_from_data(
            ['序号', '认证名称', '证书编号', '有效期', '备注'],
            [
                ['1', '【认证一】', '【证书编号】', '【有效期】', '见附件'],
                ['2', '【认证二】', '【证书编号】', '【有效期】', '见附件'],
                ['3', '【认证三】', '【证书编号】', '【有效期】', '见附件'],
                ['4', '【认证四】', '【证书编号】', '【有效期】', '见附件'],
            ]
        )
        self.add_body_paragraph('注：以上认证证书复印件加盖公章请附后，原件备查。')
        self.add_body_paragraph('（附：以上各项证明材料的复印件/扫描件）')

    def generate_section7_basic_info(self):
        """Generate Section 7: 供应商基本情况信息表."""
        self.add_page_break()
        self.add_heading1('七、供应商基本情况信息表')
        self.add_body_paragraph('按招标文件"附件1：供应商基本情况表"格式填写，可直接使用附件1模板。以下为信息汇总：')
        self.add_table_from_data(
            ['项目', '内容'],
            [
                ['供应商名称', '【投标人全称】'],
                ['统一社会信用代码', '【统一社会信用代码】'],
                ['法定代表人', '【法定代表人姓名】'],
                ['注册资本', '【注册资本】万元'],
                ['成立日期', '【成立日期】'],
                ['注册地址', '【注册地址】'],
                ['经营范围', '【经营范围】'],
                ['企业性质', '【国有企业/民营企业/外资企业/其他】'],
                ['联系人', '【联系人姓名】'],
                ['联系电话', '【联系电话】'],
                ['电子邮箱', '【电子邮箱】'],
                ['近三年营业额', '【2024年：XX万元；2025年：XX万元；2026年（预估）：XX万元】'],
                ['员工总人数', '【总人数】人'],
                ['技术人员数量', '【技术人员数量】人'],
            ]
        )
        self.add_body_paragraph('注：建议同时提交加盖公章的附件1模板版本，以符合招标文件格式要求。')

    def generate_section8_other_materials(self):
        """Generate Section 8: 采购人要求的其他材料."""
        self.add_page_break()
        self.add_heading1('八、采购人要求的其他材料')

        self.add_heading2('8.1 诚信承诺函')
        self.add_body_paragraph('诚信承诺函', bold=True)
        self.add_body_paragraph('致：【采购单位名称】')
        self.add_body_paragraph('我单位（【投标人全称】）郑重承诺：')
        self.add_body_paragraph('一、在本项目投标过程中，所提交的全部资料和信息均真实、准确、完整、有效，不存在任何虚假记载、误导性陈述或者重大遗漏。')
        self.add_body_paragraph('二、我方不存在与其他投标人串通投标的行为，也不存在以他人名义投标或者以其他方式弄虚作假骗取中标的行为。')
        self.add_body_paragraph('三、我方承诺在投标有效期内不撤销或修改投标文件。')
        self.add_body_paragraph('四、如我方中标，将严格按照招标文件要求和投标文件承诺的内容、标准提供相关服务，不将中标项目转包或违法分包。')
        self.add_body_paragraph('五、我方充分理解并同意：若上述承诺存在任何不实之处，贵方有权取消我方投标资格或中标资格，我方将依法承担相应法律责任。')
        self.add_body_paragraph('')
        self.add_body_paragraph('承诺人（公章）：【投标人全称】')
        self.add_body_paragraph('法定代表人（签字）：_______________')
        self.add_body_paragraph('日期：【    年    月    日】')

        self.add_heading2('8.2 近三年无重大违法记录声明')
        self.add_body_paragraph('我单位（【投标人全称】）郑重声明：在参加本次政府采购活动前三年内，在经营活动中没有重大违法记录，未被列入失信被执行人、重大税收违法案件当事人名单、政府采购严重违法失信行为记录名单。')
        self.add_body_paragraph('（附：信用中国网站查询截图）')

        self.add_heading2('8.3 关于非联合体投标及关联关系的声明')
        self.add_body_paragraph('我单位（【投标人全称】）郑重声明：')
        self.add_body_paragraph('一、我单位以独立投标人身份参与本项目的投标活动，未组成联合体参加本项目投标。')
        self.add_body_paragraph('二、我单位与其他参加本项目的投标单位不存在单位负责人为同一人或存在直接控股、管理关系的情形。')
        self.add_body_paragraph('三、我单位未为本采购项目提供整体设计、规范编制或者项目管理、监理、检测等服务。')
        self.add_body_paragraph('特此声明。')
        self.add_body_paragraph('')
        self.add_body_paragraph('投标人（公章）：【投标人全称】')
        self.add_body_paragraph('日期：【    年    月    日】')

    def generate_full_template(self, output_path):
        """Generate the complete bid template."""
        self.setup_page()
        self.setup_styles()

        self.generate_cover_page()
        self.generate_toc()
        self.generate_section1_proposal()
        self.generate_section2_qualifications()
        self.generate_section3_legal_rep_id()
        self.generate_section4_authorization()
        self.generate_section5_commitment()
        self.generate_section6_service_evidence()
        self.generate_section7_basic_info()
        self.generate_section8_other_materials()

        self.doc.save(output_path)
        print(f'Template generated: {output_path}')


def main():
    parser = argparse.ArgumentParser(description='Generate a bid document template')
    parser.add_argument('--output', '-o', required=True, help='Output .docx file path')
    args = parser.parse_args()

    generator = BidTemplateGenerator()
    generator.generate_full_template(args.output)


if __name__ == '__main__':
    main()

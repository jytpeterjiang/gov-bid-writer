# -*- coding: utf-8 -*-
# 盐田区政务微信授权、智能网关及相关服务保障项目 —— 响应文件格式完整版。
# 按用户提供的《第八章 响应文件格式》图片与《FB人工智能大模型管理服务第一标段》范本重构：
#   1. 增加正式响应文件前件：封面、投标函、开标一览表、服务费用说明一览表、产品清单一览表、
#      服务条款响应偏离表、授权委托书、法定代表人身份证明。
#   2. 技术方案每个小点下扩写多段详细文字，避免单段概述。
#   3. 在关键技术点插入合理示意图（总体架构、组织体系、流程图、安全体系等）。
#   4. 标题/表格/版式继续对齐 FB 范本：宋体加粗黑色、表格宋体 12pt、单线边框、1.25in 左右边距。
# 内容引号统一使用《》书名号，不在字符串内使用 ASCII 直引号。

import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

WS = os.path.dirname(os.path.abspath(__file__))
if WS not in sys.path:
    sys.path.insert(0, WS)

from generate_yantian_fb import FbBidGenerator

# —— 图片生成辅助（使用宋体，保证中文可显示）——
FONT_PATH = 'C:/Windows/Fonts/simsun.ttc'

def get_font(size):
    return ImageFont.truetype(FONT_PATH, size)


def _text_size(draw, text, font):
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_text(text, font, max_width, draw):
    """按像素宽度简单折行，支持 CJK。"""
    lines = []
    line = ""
    for char in text:
        test = line + char
        w, _ = _text_size(draw, test, font)
        if w > max_width and line:
            lines.append(line)
            line = char
        else:
            line = test
    if line:
        lines.append(line)
    return lines


def draw_box(draw, xy, text, font, fill="#E8F4F8", outline="#2E75B6", text_color="#000000"):
    """绘制带文字矩形框。"""
    x, y, x2, y2 = xy
    draw.rectangle(xy, fill=fill, outline=outline, width=2)
    w, h = x2 - x, y2 - y
    lines = wrap_text(text, font, w - 16, draw)
    line_h = font.size + 4
    total_h = len(lines) * line_h
    start_y = y + (h - total_h) / 2
    for i, ln in enumerate(lines):
        tw, _ = _text_size(draw, ln, font)
        draw.text((x + (w - tw) / 2, start_y + i * line_h), ln, font=font, fill=text_color)


def draw_arrow(draw, start, end, color="#2E75B6", width=2):
    draw.line([start, end], fill=color, width=width)
    # 简单箭头
    import math
    x0, y0 = start
    x1, y1 = end
    angle = math.atan2(y1 - y0, x1 - x0)
    arr_len = 8
    for a in (angle + 2.5, angle - 2.5):
        ax = x1 - arr_len * math.cos(a)
        ay = y1 - arr_len * math.sin(a)
        draw.line([(x1, y1), (ax, ay)], fill=color, width=width)


def create_architecture_diagram(path):
    """图1：系统总体架构图。"""
    W, H = 700, 460
    img = Image.new('RGB', (W, H), 'white')
    draw = ImageDraw.Draw(img)
    font = get_font(18)
    small = get_font(14)
    layers = [
        ("终端用户层", 50, 30, 650, 80, "PC/移动端/信创终端/微信小程序", "#D5E8F0"),
        ("接入与网关层", 50, 120, 650, 170, "智能网关/负载均衡/SSL 卸载/访问控制", "#E1F0E1"),
        ("平台服务层", 50, 210, 650, 260, "政务微信/消息推送/应用中心/直播/会议", "#FFF2CC"),
        ("数据与接口层", 50, 300, 650, 350, "组织/消息/日志/统计/第三方业务接口", "#FCE4D6"),
        ("运维保障体系", 50, 390, 650, 440, "监控/告警/SLA/安全/培训/应急响应", "#E2E0F9"),
    ]
    for name, x1, y1, x2, y2, sub, fill in layers:
        draw_box(draw, (x1, y1, x2, y2), name, font, fill=fill)
        tw, _ = _text_size(draw, sub, small)
        draw.text((x1 + (x2 - x1 - tw) / 2, y2 + 3), sub, font=small, fill="#555555")
    for i in range(len(layers) - 1):
        draw_arrow(draw, (W / 2, layers[i][4] + 5), (W / 2, layers[i + 1][1] - 5))
    img.save(path)


def create_org_chart(path):
    """图2：项目服务组织架构图。"""
    W, H = 700, 400
    img = Image.new('RGB', (W, H), 'white')
    draw = ImageDraw.Draw(img)
    font = get_font(16)
    draw_box(draw, (250, 20, 450, 70), "项目总监", font, fill="#FCE4D6")
    draw_arrow(draw, (350, 70), (350, 100))
    draw_box(draw, (250, 100, 450, 150), "项目经理", font, fill="#D5E8F0")
    # 下挂四个组
    groups = [
        (50, 190, 180, 240, "技术运维组"),
        (200, 190, 330, 240, "安全合规组"),
        (370, 190, 500, 240, "运营服务组"),
        (540, 190, 670, 240, "应急响应组"),
    ]
    cx = 350
    for x1, y1, x2, y2, name in groups:
        draw_arrow(draw, (cx, 150), (cx, 100))
        draw_arrow(draw, (cx, 150), (x1 + (x2 - x1) / 2, y1))
        draw_box(draw, (x1, y1, x2, y2), name, font, fill="#E2E0F9")
    # 各组下职责
    duties = [
        ("网关/平台/应用运维", 50, 260, 180),
        ("等保/保密/审计", 200, 260, 330),
        ("用户/数据/应用运营", 370, 260, 500),
        ("7x24 值班/演练", 540, 260, 670),
    ]
    small = get_font(13)
    for txt, x1, y, x2 in duties:
        tw, _ = _text_size(draw, txt, small)
        draw.text((x1 + (x2 - x1 - tw) / 2, y), txt, font=small, fill="#333333")
    img.save(path)


def create_service_process_diagram(path):
    """图3：ITIL 服务流程图。"""
    W, H = 800, 220
    img = Image.new('RGB', (W, H), 'white')
    draw = ImageDraw.Draw(img)
    font = get_font(15)
    boxes = [
        (30, 60, 150, 110, "事件受理"),
        (180, 60, 300, 110, "分类定级"),
        (330, 60, 450, 110, "诊断处置"),
        (480, 60, 600, 110, "恢复验证"),
        (630, 60, 750, 110, "闭环归档"),
    ]
    for x1, y1, x2, y2, name in boxes:
        draw_box(draw, (x1, y1, x2, y2), name, font, fill="#E1F0E1")
    for i in range(len(boxes) - 1):
        draw_arrow(draw, (boxes[i][2] + 5, 85), (boxes[i + 1][1] - 5, 85))
    # 下方反馈
    draw_box(draw, (280, 140, 520, 190), "持续改进 / 知识沉淀 / 月度复盘", font, fill="#FFF2CC")
    draw_arrow(draw, (400, 110), (400, 140))
    img.save(path)


def create_sla_escalation_diagram(path):
    """图4：SLA 分级响应与升级路径。"""
    W, H = 700, 300
    img = Image.new('RGB', (W, H), 'white')
    draw = ImageDraw.Draw(img)
    font = get_font(15)
    levels = [
        (50, 50, 200, 100, "一级\n核心系统中断", "15 分钟", "2 小时"),
        (270, 50, 420, 100, "二级\n重要功能异常", "30 分钟", "4 小时"),
        (490, 50, 640, 100, "三级\n一般问题", "60 分钟", "8 小时"),
    ]
    for x1, y1, x2, y2, title, resp, recov in levels:
        draw_box(draw, (x1, y1, x2, y2), title, font, fill="#FCE4D6")
        small = get_font(13)
        line = f"响应 {resp} | 恢复 {recov}"
        tw, _ = _text_size(draw, line, small)
        draw.text((x1 + (x2 - x1 - tw) / 2, y2 + 8), line, font=small, fill="#333333")
    # 升级路径
    draw_arrow(draw, (125, 130), (350, 180))
    draw_arrow(draw, (345, 130), (350, 180))
    draw_arrow(draw, (565, 130), (350, 180))
    draw_box(draw, (220, 180, 480, 230), "二线专家 / 原厂支持 / 管理升级", font, fill="#E2E0F9")
    draw_arrow(draw, (350, 230), (350, 270))
    draw_box(draw, (220, 270, 480, 295), "项目总监/采购人决策层", get_font(13), fill="#D5E8F0")
    img.save(path)


def create_security_architecture_diagram(path):
    """图5：安全保障体系架构图。"""
    W, H = 700, 360
    img = Image.new('RGB', (W, H), 'white')
    draw = ImageDraw.Draw(img)
    font = get_font(16)
    layers = [
        ("安全管理制度", 50, 30, 650, 80, "#E2E0F9"),
        ("安全组织与人员", 50, 110, 650, 160, "#D5E8F0"),
        ("安全技术控制", 50, 190, 650, 240, "#FFF2CC"),
        ("安全运维与审计", 50, 270, 650, 320, "#FCE4D6"),
    ]
    for name, x1, y1, x2, y2, fill in layers:
        draw_box(draw, (x1, y1, x2, y2), name, font, fill=fill)
    for i in range(len(layers) - 1):
        draw_arrow(draw, (W / 2, layers[i][4] + 5), (W / 2, layers[i + 1][1] - 5))
    # 横向控制点
    controls = ["身份鉴别", "访问控制", "传输加密", "日志审计", "漏洞管理", "备份恢复"]
    small = get_font(13)
    cw = (650 - 50) / len(controls)
    for i, c in enumerate(controls):
        x = 50 + i * cw + 10
        draw_box(draw, (x, 250, x + cw - 20, 280), c, small, fill="#E1F0E1", outline="#70AD47")
    img.save(path)


def create_emergency_response_diagram(path):
    """图6：应急响应流程图。"""
    W, H = 700, 320
    img = Image.new('RGB', (W, H), 'white')
    draw = ImageDraw.Draw(img)
    font = get_font(15)
    boxes = [
        (250, 20, 450, 70, "事件发现与告警"),
        (250, 100, 450, 150, "分级定级与通报"),
        (100, 180, 300, 230, "技术处置"),
        (400, 180, 600, 230, "业务恢复"),
        (250, 260, 450, 310, "复盘改进与报告"),
    ]
    for x1, y1, x2, y2, name in boxes:
        draw_box(draw, (x1, y1, x2, y2), name, font, fill="#D5E8F0")
    # arrows
    draw_arrow(draw, (350, 70), (350, 100))
    draw_arrow(draw, (250, 150), (200, 180))
    draw_arrow(draw, (450, 150), (500, 180))
    draw_arrow(draw, (200, 230), (350, 260))
    draw_arrow(draw, (500, 230), (450, 260))
    img.save(path)


def create_training_system_diagram(path):
    """图7：培训与知识转移体系。"""
    W, H = 700, 300
    img = Image.new('RGB', (W, H), 'white')
    draw = ImageDraw.Draw(img)
    font = get_font(16)
    draw_box(draw, (50, 30, 650, 80), "培训需求分析", font, fill="#E2E0F9")
    draw_arrow(draw, (350, 80), (350, 110))
    modules = [
        (50, 110, 200, 160, "管理员培训"),
        (270, 110, 420, 160, "运维培训"),
        (490, 110, 640, 160, "最终用户培训"),
    ]
    for x1, y1, x2, y2, name in modules:
        draw_box(draw, (x1, y1, x2, y2), name, font, fill="#D5E8F0")
    draw_arrow(draw, (350, 160), (350, 190))
    draw_box(draw, (50, 190, 650, 240), "考核认证 + 知识库 + 持续辅导", font, fill="#FFF2CC")
    draw_arrow(draw, (350, 240), (350, 270))
    draw_box(draw, (50, 270, 650, 295), "自主运营能力移交", get_font(14), fill="#E1F0E1")
    img.save(path)


def create_implementation_roadmap(path):
    """图8：项目实施路线图。"""
    W, H = 800, 220
    img = Image.new('RGB', (W, H), 'white')
    draw = ImageDraw.Draw(img)
    font = get_font(14)
    phases = [
        (30, 60, 150, 110, "第1-2月\n进场与盘点"),
        (170, 60, 290, 110, "第3-4月\n体系搭建"),
        (310, 60, 430, 110, "第5-6月\n首轮优化"),
        (450, 60, 570, 110, "第7-9月\n稳定运营"),
        (590, 60, 710, 110, "第10-12月\n验收移交"),
    ]
    for x1, y1, x2, y2, name in phases:
        draw_box(draw, (x1, y1, x2, y2), name, font, fill="#E1F0E1")
    for i in range(len(phases) - 1):
        draw_arrow(draw, (phases[i][2] + 5, 85), (phases[i + 1][1] - 5, 85))
    # milestone bar
    milestones = [
        (90, 150, "基线报告"),
        (230, 150, "制度发布"),
        (370, 150, "首季报告"),
        (510, 150, "合规报告"),
        (650, 150, "终验总结"),
    ]
    small = get_font(13)
    for x, y, txt in milestones:
        tw, _ = _text_size(draw, txt, small)
        draw.text((x - tw / 2, y), txt, font=small, fill="#333333")
    img.save(path)


class ResponseBidGenerator(FbBidGenerator):
    """
    响应文件格式完整版生成器。
    在 FbBidGenerator 基础上：
      - 增加正式响应文件前件；
      - 技术方案各小点扩写详细文字；
      - 插入关键技术示意图。
    """

    def __init__(self):
        super().__init__()
        self.image_dir = os.path.join(WS, 'response_images')
        os.makedirs(self.image_dir, exist_ok=True)
        self.generate_all_images()

    def generate_all_images(self):
        """预生成所有示意图。"""
        funcs = [
            create_architecture_diagram,
            create_org_chart,
            create_service_process_diagram,
            create_sla_escalation_diagram,
            create_security_architecture_diagram,
            create_emergency_response_diagram,
            create_training_system_diagram,
            create_implementation_roadmap,
        ]
        for idx, fn in enumerate(funcs, 1):
            fn(os.path.join(self.image_dir, f"diagram_{idx:02d}.png"))

    def _image_path(self, idx):
        return os.path.join(self.image_dir, f"diagram_{idx:02d}.png")

    def _add_picture(self, caption, idx, width=Inches(6.0)):
        """插入图片并加图题。"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(self._image_path(idx), width=width)
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"图 {idx}  {caption}")
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(10.5)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # —— 正式响应文件前件 ——
    def generate_cover_page(self):
        """响应文件封面，按图片格式。"""
        # 正副本标识
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run('【正/副本】')
        r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(14); r.font.bold = True

        # 空行
        for _ in range(6):
            self.doc.add_paragraph()

        # 项目名称
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('盐田区政务微信授权、智能网关及相关服务保障项目')
        r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(22); r.font.bold = True

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('项 目 响 应 文 件')
        r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(26); r.font.bold = True

        for _ in range(8):
            self.doc.add_paragraph()

        # 底部信息
        items = [
            '项目编号：                  （待填写）',
            '投标人（全称）：            （盖章）',
            '法定代表人或其授权委托人：  （签字或盖章）',
            '地    址：',
            '联 系 人：                  联系电话：',
            '日    期：      年      月      日',
        ]
        for it in items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(it)
            r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(14)
        self.doc.add_page_break()

    def _add_center_title(self, text, size=Pt(16), bold=True):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = size; r.font.bold = bold
        return p

    def _add_plain_paragraph(self, text, indent=0.0, first_line=0.0):
        p = self.doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        if first_line:
            p.paragraph_format.first_line_indent = Cm(first_line)
        r = p.add_run(text)
        r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(12)
        return p

    def _add_bold_field(self, label, value=''):
        p = self.doc.add_paragraph()
        r1 = p.add_run(label)
        r1.font.name = '宋体'; r1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r1.font.size = Pt(12); r1.font.bold = True
        if value:
            r2 = p.add_run(value)
            r2.font.name = '宋体'; r2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r2.font.size = Pt(12)
        return p

    def generate_bid_letter(self):
        """投标函。"""
        self._add_center_title('投 标 函', Pt(18), True)
        self._add_plain_paragraph('（招标人名称）：')
        self._add_plain_paragraph(
            '我方收到贵方关于《盐田区政务微信授权、智能网关及相关服务保障项目》的招标文件，'
            '经认真研究，愿意按照招标文件中的一切要求提供相应服务，完成合同约定的责任和义务。'
            '现就有关事项郑重承诺如下：'
        )
        letter_items = [
            '我方愿意按照招标文件中的一切要求，提供盐田区政务微信授权、智能网关及相关服务保障项目的全部服务内容，'
            '完成合同约定的责任和义务，确保服务质量达到现行合格标准。',
            '我方已详细阅读并完全理解本项目招标文件、补充通知、答疑文件及评标办法的全部内容，'
            '同意放弃对招标文件中任何含糊不清或误解问题提出异议的权利。',
            '如果我方在投标有效期内撤回响应文件或承诺，愿承担由此产生的一切法律责任。',
            '我方同意向贵方提供可能要求的、与本次投标有关的任何证据资料，并保证所提供资料真实、完整、有效。',
            '我方的响应文件自响应文件递交截止之日起九十天内有效。',
            '如果我方中标，将严格按照招标文件要求完成全部内容，服务质量达到现行合格标准，并符合国家、行业、地方及招标文件的相关规定。',
            '我方完全接受并响应招标文件、答疑文件、评标办法、采购预算及限价等关于本项目相关文件的要求，'
            '严格遵守开标过程的时间安排及各项纪律。',
            '所有关于本响应文件的函电，请按本响应文件投标人信息部分所列地址和联系方式联系。',
        ]
        for i, it in enumerate(letter_items, 1):
            self._add_plain_paragraph(f'{i}. {it}', first_line=0.0)
        self._add_plain_paragraph('')
        self._add_plain_paragraph('投标人：                    （盖章）')
        self._add_plain_paragraph('法定代表人或被授权人：      （签字或盖章）')
        self._add_plain_paragraph('地    址：')
        self._add_plain_paragraph('开户银行：')
        self._add_plain_paragraph('账    号：')
        self._add_plain_paragraph('电    话：')
        self._add_plain_paragraph('日    期：      年      月      日')
        self.doc.add_page_break()

    def _make_simple_table(self, headers, rows, col_widths=None):
        """创建简单单线表格，表头浅蓝底。"""
        from docx.oxml import parse_xml
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            for p in hdr[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    r.font.size = Pt(12); r.font.bold = True
            # 浅蓝底纹
            tcPr = hdr[i]._tc.get_or_add_tcPr()
            tcPr.append(parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D5E8F0"/>'))
        for row in rows:
            cells = table.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = str(v)
                for p in cells[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        r.font.size = Pt(12)
        return table

    def generate_opening_table(self):
        """开标一览表。"""
        self._add_center_title('开 标 一 览 表', Pt(18), True)
        self._add_plain_paragraph('项目名称：盐田区政务微信授权、智能网关及相关服务保障项目')
        self._add_plain_paragraph('项目编号：（待填写）')
        self._add_plain_paragraph('项目标段：/')
        self._add_plain_paragraph('')
        self._make_simple_table(
            ['序号', '服务内容', '投标报价（元）', '服务期', '备注'],
            [
                ['1', '政务微信订阅许可续约及相关运营服务', '', '一年', ''],
                ['2', '智能网关维保及相关服务保障', '', '一年', ''],
                ['3', '合计总价', '', '一年', ''],
            ]
        )
        self._add_plain_paragraph('')
        self._add_plain_paragraph('投标人：                    （盖章）')
        self._add_plain_paragraph('法定代表人或被授权人：      （签字或盖章）')
        self._add_plain_paragraph('日    期：      年      月      日')
        self.doc.add_page_break()

    def generate_cost_explanation_table(self):
        """服务费用说明一览表。"""
        self._add_center_title('服 务 费 用 说 明 一 览 表', Pt(18), True)
        self._add_plain_paragraph('项目名称：盐田区政务微信授权、智能网关及相关服务保障项目')
        self._add_plain_paragraph('项目编号：（待填写）')
        self._add_plain_paragraph('')
        self._add_plain_paragraph('说明：')
        self._add_plain_paragraph('1. 所有投标报价均含税，且用人民币表示，单位为元，精确到小数点后两位。', first_line=0.74)
        self._add_plain_paragraph('2. 该表中包含投标人认为完成本项目所需的所有费用，合计总价应与开标一览表中投标总报价金额一致。', first_line=0.74)
        self._add_plain_paragraph('3. 报价分项以招标文件要求的服务目录为基础，确保无隐性费用、无二次收费。', first_line=0.74)
        self._add_plain_paragraph('')
        self._make_simple_table(
            ['序号', '费用项目', '说明', '金额（元）'],
            [
                ['1', '软件许可与订阅费用', '政务微信订阅席位、软件使用许可、云服务调用额度等', ''],
                ['2', '运维服务费用', '远程支持、现场服务、健康检查、应急响应、运行保障等', ''],
                ['3', '运营服务费用', '组织和用户运营、应用运营、数据运营统计等', ''],
                ['4', '专家咨询与培训费用', '专家评审、培训教材、现场授课、考核认证等', ''],
                ['5', '信创适配与升级费用', '信创终端适配、版本升级、兼容性验证等', ''],
                ['6', '其他费用', '投标人认为完成项目所需的其他必要费用', ''],
                ['7', '合计', '合计总价（应与开标一览表一致）', ''],
            ]
        )
        self._add_plain_paragraph('')
        self._add_plain_paragraph('投标人名称：                （盖章）')
        self._add_plain_paragraph('法定代表人或被授权人：      （签字或盖章）')
        self._add_plain_paragraph('日    期：      年      月      日')
        self.doc.add_page_break()

    def generate_product_list_table(self):
        """产品清单一览表。"""
        self._add_center_title('产 品 清 单 一 览 表', Pt(18), True)
        self._add_plain_paragraph('备注：服务所用产品需提供产品厂商及型号、功能、参数等信息。')
        self._add_plain_paragraph('')
        self._make_simple_table(
            ['序号', '产品/服务名称', '厂商', '型号/版本', '功能简述', '数量', '单位'],
            [
                ['1', '政务微信协同平台订阅许可', '腾讯', '政务微信', '提供即时通讯、组织管理、应用集成等能力', '4300', '席/年'],
                ['2', '智能网关服务', '供应商', '定制', '提供访问控制、链路转发、协议适配、日志审计等', '1', '项/年'],
                ['3', '运维管理平台', '供应商', '定制', '工单、监控、报表、知识库、CMDB 等', '1', '套/年'],
                ['4', '远程支持服务', '供应商', '服务', '7x24 热线、远程诊断、远程配置', '1', '项/年'],
                ['5', '培训服务', '供应商', '服务', '管理员、运维、最终用户培训与考核', '1', '项/年'],
            ]
        )
        self.doc.add_page_break()

    def generate_deviation_table(self):
        """服务条款响应偏离表。"""
        self._add_center_title('服 务 条 款 响 应 偏 离 表', Pt(18), True)
        self._add_plain_paragraph('项目名称：盐田区政务微信授权、智能网关及相关服务保障项目')
        self._add_plain_paragraph('项目编号：（待填写）')
        self._add_plain_paragraph('')
        self._make_simple_table(
            ['序号', '招标文件条款', '投标文件响应', '偏离情况', '说明'],
            [
                ['1', '服务期限：一年', '服务期限：一年', '无偏离', ''],
                ['2', '服务响应时间要求', '按一级/二级/三级分级响应', '无偏离', '详见 SLA 章节'],
                ['3', '服务团队人员要求', '配置项目经理、技术运维、安全合规、运营服务、应急响应等人员', '无偏离', ''],
                ['4', '保密与信息安全要求', '签署保密协议、按等保 2.0 要求执行、全程审计', '无偏离', ''],
                ['5', '培训与知识转移要求', '提供管理员、运维、用户培训，考核认证', '无偏离', ''],
                ['6', '验收与考核要求', '按月/季/年提交报告，配合满意度调查与终验', '无偏离', ''],
            ]
        )
        self._add_plain_paragraph('')
        self._add_plain_paragraph('投标人名称：                （盖章）')
        self._add_plain_paragraph('法定代表人或被授权人：      （签字或盖章）')
        self._add_plain_paragraph('日    期：      年      月      日')
        self.doc.add_page_break()

    def generate_authorization_letter(self):
        """授权委托书。"""
        self._add_center_title('授 权 委 托 书', Pt(18), True)
        self._add_plain_paragraph('')
        self._add_plain_paragraph('本人        （姓名）系        （供应商名称）的法定代表人，'
                                  '现委托        （姓名）为我方代理人。')
        self._add_plain_paragraph(
            '代理人根据授权，以我方名义签署、澄清、说明、补正、递交、撤回、修改'
            '《盐田区政务微信授权、智能网关及相关服务保障项目》项目响应文件、'
            '签订合同和处理有关事宜，其法律后果由我方承担。'
        )
        self._add_plain_paragraph('')
        self._add_plain_paragraph('委托期限：')
        self._add_plain_paragraph('')
        self._add_plain_paragraph('代理人无转委托权。')
        self._add_plain_paragraph('')
        self._add_plain_paragraph('附：法定代表人身份证明、委托代理人身份证明')
        self._add_plain_paragraph('')
        self._add_plain_paragraph('投标人：                    （盖章）')
        self._add_plain_paragraph('法定代表人：                （签字或盖章）')
        self._add_plain_paragraph('身份证号码：')
        self._add_plain_paragraph('委托代理人：                （签字或盖章）')
        self._add_plain_paragraph('身份证号码：')
        self._add_plain_paragraph('日    期：      年      月      日')
        self.doc.add_page_break()

    def generate_legal_id(self):
        """法定代表人身份证明。"""
        self._add_center_title('法 定 代 表 人 身 份 证 明', Pt(18), True)
        self._add_plain_paragraph('')
        self._add_plain_paragraph('投标人名称：')
        self._add_plain_paragraph('单位性质：')
        self._add_plain_paragraph('地    址：')
        self._add_plain_paragraph('')
        self._add_plain_paragraph(
            '        姓名：        性别：        年龄：        职务：        '
            '系        （投标人名称）的法定代表人。'
        )
        self._add_plain_paragraph('')
        self._add_plain_paragraph('特此证明。')
        self._add_plain_paragraph('')
        self._add_plain_paragraph('')
        self._add_plain_paragraph('投标人：                    （盖章）')
        self._add_plain_paragraph('日    期：      年      月      日')
        self.doc.add_page_break()

    def generate_front_matter(self):
        """生成完整响应文件前件。"""
        self.generate_cover_page()
        self.generate_bid_letter()
        self.generate_opening_table()
        self.generate_cost_explanation_table()
        self.generate_product_list_table()
        self.generate_deviation_table()
        self.generate_authorization_letter()
        self.generate_legal_id()

    # —— 技术方案（扩写 + 图片）——
    def generate_technical_proposal(self):
        """重写技术方案：每个小点多段文字，并插入关键示意图。"""
        self.add_heading1('第一章 技术方案概述与总体设计')
        self.add_body_paragraph(
            '本技术标针对《盐田区政务微信授权、智能网关及相关服务保障项目》的采购需求，'
            '系统阐述我方对项目背景、服务范围、总体架构、实施方案、运维保障、质量管理、'
            '信息安全、应急响应、培训验收等方面的理解、承诺与具体措施。以下各节将围绕招标文件中的关键要求，'
            '以逐条响应、分层展开、图文并茂的方式，全面展现我方的服务能力与落地路径。'
        )

        # 1.1 项目背景与理解
        self.add_heading2('1.1 项目背景与需求理解')
        self.add_body_paragraph(
            '盐田区作为深圳市重要的政务协同先行区域，对政务微信、智能网关等协同平台的稳定性、安全性、'
            '可用性提出了持续性要求。本项目涉及订阅许可续约、信创环境适配、软件使用许可、新版本升级许可、'
            '云服务能力许可、售后响应、远程授权、远程技术支持、组织和用户运营、数据运营统计、应用运营、'
            '运行保障、专家咨询以及智能网关维保等多项服务内容，具有服务对象广、服务周期紧、合规要求严、'
            '技术耦合度高等特点。'
        )
        self.add_body_paragraph(
            '我方理解，本项目不仅是软件许可与运维服务的采购，更是对一套完整、可持续、可度量的政务协同'
            '服务保障体系的建设要求。采购人期望通过本次采购，实现平台可用性不低于既定目标、故障响应'
            '及时有效、用户满意度持续提升、安全合规全面达标、知识资产有序沉淀等多重目标。因此，'
            '我方的方案设计以 ITIL 4 服务管理框架为方法论基础，以 ISO/IEC 27001:2022 信息安全管理体系为'
            '安全合规底座，以等保 2.0 为技术控制标尺，形成覆盖人员、流程、技术、制度四维一体的服务方案。'
        )
        self.add_body_paragraph(
            '在需求理解层面，我方将服务需求归纳为以下五大类：一是许可与版本保障类，确保席位、软件、'
            '云服务、新版本等授权在合同期内持续有效；二是运维与响应保障类，提供 7x24 受理、分级响应、'
            '现场与远程支持；三是运营与数据支撑类，涵盖组织治理、应用运营、数据统计与洞察；四是安全与'
            '合规保障类，覆盖等保、保密、账号权限、日志审计；五是持续改进与知识转移类，通过培训、复盘、'
            '知识库、SOP 优化实现服务螺旋上升。以下各节将逐一展开。'
        )

        # 1.2 总体技术方案
        self.add_heading2('1.2 总体技术方案与架构设计')
        self.add_body_paragraph(
            '我方总体方案采用分层解耦、冗余可靠、安全可控的设计原则。整体架构自下而上分为终端用户层、'
            '接入与网关层、平台服务层、数据与接口层、运维保障层五个层次。各层之间通过标准化接口进行'
            '数据交换，通过统一身份认证与访问控制实现安全边界，通过智能网关实现内外网访问的收敛与转发，'
            '通过运维保障体系实现全生命周期的监控、响应、优化与审计。'
        )
        self.add_body_paragraph(
            '终端用户层覆盖 PC 客户端、移动端、信创终端以及微信小程序等多种访问形态，满足不同岗位、'
            '不同场景的用户需求。接入与网关层承担 SSL 卸载、协议转换、访问控制、流量调度、日志留存的'
            '关键职责，是内外网交互的安全闸门。平台服务层提供政务微信核心能力，包括即时通讯、组织通讯录、'
            '应用中心、消息推送、直播会议、审批流转等。数据与接口层负责与采购人现有业务系统的对接，'
            '包括组织架构同步、单点登录、消息回执、数据看板等。运维保障层贯穿服务全过程，提供监控告警、'
            '事件管理、问题管理、变更管理、配置管理、知识管理、安全审计等能力。'
        )
        self.add_body_paragraph(
            '在架构实现上，我方强调四个关键特性：一是高可用性，通过双链路冗余、网关主备、负载均衡等'
            '机制降低单点故障风险；二是可扩展性，平台服务层支持按组织规模、业务场景动态扩展；三是'
            '安全合规性，通过最小权限、加密传输、访问控制、日志审计等措施满足等保 2.0 要求；四是'
            '可运维性，通过标准化 SOP、自动化监控、统一运维平台提升服务效率与可追溯性。图 1 展示了'
            '系统总体架构。'
        )
        self._add_picture('系统总体架构图', 1, width=Inches(6.2))

        self.add_heading3('1.2.1 智能网关设计方案')
        self.add_body_paragraph(
            '智能网关是本项目的关键基础设施，承担着外部访问收敛、协议适配、安全控制、流量调度等核心职责。'
            '我方设计的网关方案包含接入网关、协议网关、安全网关、管理网关四个逻辑组件。接入网关负责'
            'HTTPS 请求的 SSL 卸载与会话保持；协议网关负责将外部请求转换为内部服务可识别的协议格式；'
            '安全网关负责身份校验、访问控制、威胁检测、日志留存；管理网关负责配置下发、灰度发布、'
            '健康检查与策略管理。'
        )
        self.add_body_paragraph(
            '在具体部署上，我方建议采用主备双活模式，主节点承担日常流量，备节点实时同步配置与状态，'
            '当主节点发生故障时可在分钟级完成切换。网关的配置管理纳入 CMDB，所有变更均通过变更管理'
            '流程审批并留痕。证书管理采用自动到期预警机制，防止因证书过期导致服务中断。日志留存'
            '不少于六个月，并支持按用户、时间、接口、结果等多维度检索。'
        )

        self.add_heading3('1.2.2 信创环境适配方案')
        self.add_body_paragraph(
            '信创环境适配是本项目的重要内容。我方将信创适配分为终端适配、操作系统适配、浏览器内核适配、'
            '外设适配四个层面。终端适配主要覆盖鲲鹏、飞腾等主流信创 CPU 架构；操作系统适配覆盖统信 UOS、'
            '麒麟操作系统等；浏览器内核适配覆盖 Chromium、WebKit 等主流内核及信创浏览器；外设适配覆盖'
            '摄像头、麦克风、打印机等常见办公外设。'
        )
        self.add_body_paragraph(
            '适配工作按发现、评估、修复、验证、上线、跟踪六步闭环执行。发现阶段通过用户反馈、兼容性'
            '测试、厂商公告等渠道收集问题；评估阶段判断影响范围、优先级与修复方案；修复阶段通过补丁、'
            '配置调整、版本升级等方式解决；验证阶段在信创测试环境中进行回归测试；上线阶段通过灰度发布'
            '降低风险；跟踪阶段通过满意度回访与问题复发率监控验证效果。所有适配案例均纳入知识库，'
            '形成可复用的经验资产。'
        )

        # 1.3 服务实施与组织保障
        self.add_heading2('1.3 服务实施与组织保障')
        self.add_body_paragraph(
            '我方将为本项目配置专业化、稳定化、协同化的服务团队。团队采用项目总监负责制，下设项目经理、'
            '技术运维组、安全合规组、运营服务组、应急响应组。项目总监负责重大事项决策、资源协调与'
            '客户关系管理；项目经理负责日常服务计划、进度跟踪、质量把控与报告输出；技术运维组负责平台、'
            '网关、应用的日常运维与故障处置；安全合规组负责等保、保密、审计、权限治理；运营服务组负责'
            '用户运营、应用运营、数据运营；应急响应组负责 7x24 值班、突发事件处置与演练。'
        )
        self.add_body_paragraph(
            '人员配置方面，我方承诺项目经理具备同类政务协同项目五年以上管理经验，技术负责人具备平台运维、'
            '网关运维、安全合规相关资质与经验，关键岗位设置 AB 角，确保人员请假、离职、突发情况下服务'
            '不中断。所有驻场人员上岗前接受保密培训与项目专项培训，签署保密协议与服务承诺书。项目团队'
            '实行月度例会、周例会、日站会三级沟通机制，确保信息同步与问题快速升级。图 2 展示了项目服务'
            '组织架构。'
        )
        self._add_picture('项目服务组织架构图', 2, width=Inches(6.2))

        self.add_heading3('1.3.1 实施路线图')
        self.add_body_paragraph(
            '项目实施按进场准备、体系搭建、首轮优化、稳定运营、验收移交五个阶段推进。进场准备阶段完成'
            '资产盘点、账号梳理、环境熟悉、工具部署与团队入驻；体系搭建阶段发布管理制度、SOP、检查单、'
            '服务目录、SLA 与 KPI 指标；首轮优化阶段完成首次全覆盖巡检、隐患清零、配置优化与基线建立；'
            '稳定运营阶段进入日常运维、事件响应、定期巡检、报告输出与持续改进；验收移交阶段完成年度'
            '总结、知识资产移交、满意度调查与终验配合。图 8 展示了项目实施路线图。'
        )
        self._add_picture('项目实施路线图', 8, width=Inches(6.2))

        # 1.4 运维保障体系
        self.add_heading2('1.4 运维保障体系与流程设计')
        self.add_body_paragraph(
            '我方以 ITIL 4 服务管理框架为方法论，建立事件管理、问题管理、变更管理、配置管理、发布管理、'
            '知识管理、服务请求管理七大流程。事件管理强调快速受理、分级响应、限时恢复、闭环归档；'
            '问题管理强调根因分析、预防措施、知识沉淀；变更管理强调审批、测试、回滚、灰度；配置管理'
            '强调 CMDB 准确、变更留痕、版本可控；发布管理强调计划、评审、测试、上线、验证；知识管理'
            '强调经验复用、案例积累、SOP 优化；服务请求管理强调标准服务、自助服务、满意度跟踪。'
        )
        self.add_body_paragraph(
            '事件管理流程是运维保障的核心。当用户或监控系统发现异常时，首先由服务台进行受理与分类，'
            '根据影响范围与紧急程度确定事件等级；随后进入诊断处置阶段，一线进行常规处置，二线处理复杂'
            '问题，必要时引入原厂或第三方支持；处置完成后进行恢复验证，确认业务功能正常；最后闭环归档，'
            '记录根因、处置过程、改进措施，并同步更新知识库。图 3 展示了 ITIL 服务流程。'
        )
        self._add_picture('ITIL 服务流程图', 3, width=Inches(6.2))

        self.add_heading3('1.4.1 分级响应与 SLA')
        self.add_body_paragraph(
            '我方将事件分为一级、二级、三级三个等级。一级事件指核心系统全面中断或主要功能大面积不可用，'
            '影响范围覆盖全区或关键业务；二级事件指重要功能异常，影响部分部门或用户；三级事件指一般'
            '性问题，影响个别用户或单一操作。对应不同等级，我方设定不同的响应时间、恢复时间与升级路径，'
            '确保资源优先投入到影响最大的事件。'
        )
        self.add_body_paragraph(
            '响应时间方面，一级事件 15 分钟内响应、2 小时内恢复；二级事件 30 分钟内响应、4 小时内恢复；'
            '三级事件 60 分钟内响应、8 小时内恢复。若超出响应或恢复时限，系统自动触发升级，由项目经理、'
            '项目总监逐级介入，必要时向采购人决策层通报。图 4 展示了 SLA 分级响应与升级路径。'
        )
        self._add_picture('SLA 分级响应与升级路径', 4, width=Inches(6.2))

        # 1.5 质量管理与考核
        self.add_heading2('1.5 质量管理与考核体系')
        self.add_body_paragraph(
            '我方以 PDCA 循环驱动服务质量持续提升。计划阶段，根据招标文件与合同约定明确 SLA、KPI、'
            '服务目录、交付物清单与验收标准；执行阶段，严格按照 SOP 与流程交付服务，每日站会、每周例会、'
            '每月复盘；检查阶段，通过运维平台、看板、月报、满意度调查、内部审计等方式度量服务效果；'
            '改进阶段，针对检查发现的问题与偏差，制定改进项、分配责任人、设定完成时限，并跟踪闭环。'
        )
        self.add_body_paragraph(
            '核心 KPI 包括响应时效达标率、故障恢复平均时长、重大故障次数、用户满意度、知识库增量、'
            '安全保密事件数等。其中响应时效达标率目标不低于 99%，故障恢复平均时长按分级要求执行，'
            '重大故障年度为零目标，用户满意度季度调研不低于 90 分，知识库每季度新增或修订不少于若干条，'
            '安全保密事件年度为零。所有 KPI 均纳入月度报告，供采购人监督与考核。'
        )
        self.add_body_paragraph(
            '考核结果与服务改进直接挂钩。若某项 KPI 未达标，我方将在 5 个工作日内提交原因分析与整改计划；'
            '若连续两个月未达标，我方将启动专项改进，由项目总监亲自督导，必要时调整人员或资源配置。'
            '考核优秀的经验将被提炼为最佳实践，纳入培训教材与知识库。'
        )

        # 1.6 信息安全与保密
        self.add_heading2('1.6 信息安全与保密保障')
        self.add_body_paragraph(
            '信息安全是本项目的重要底线。我方依据 ISO/IEC 27001:2022、网络安全等级保护 2.0 以及采购人'
            '相关保密制度，建立涵盖管理制度、安全组织、安全技术、安全运维四个层面的保障体系。'
            '管理制度层面，我方提交《信息安全与保密管理制度》《账号权限管理制度》《数据安全管理制度》等；'
            '安全组织层面，设立安全负责人与保密员，明确岗位职责与汇报路径；安全技术层面，落实身份鉴别、'
            '访问控制、传输加密、日志审计、漏洞管理、备份恢复等措施；安全运维层面，将巡检、漏洞修复、'
            '应急响应、配置审计纳入 SOP 与月度报告。'
        )
        self.add_body_paragraph(
            '在账号与权限管理方面，我方实行实名制、最小权限、定期复核、及时回收四项原则。所有账号与权限'
            '变更均通过审批流程，操作留痕，支持审计追溯。Inactive 账号每季度清理一次，离职或调岗人员权限'
            '在 24 小时内回收。在数据安全方面，重要数据访问需审批，传输采用加密通道，日志留存不少于六个月，'
            '禁止未经授权的数据复制、导出与外发。图 5 展示了安全保障体系架构。'
        )
        self._add_picture('安全保障体系架构图', 5, width=Inches(6.2))

        # 1.7 应急响应与业务连续性
        self.add_heading2('1.7 应急响应与业务连续性保障')
        self.add_body_paragraph(
            '我方建立三级应急响应体系：一线值守负责日常事件受理与初级处置；二线专家负责复杂技术问题诊断'
            '与恢复；三线支持包括原厂、运营商、第三方供应商等外部资源。应急响应组实行 7x24 值班制度，'
            '值班人员携带专用应急手机，确保任何时间均可响应。'
        )
        self.add_body_paragraph(
            '应急响应流程分为事件发现与告警、分级定级与通报、技术处置、业务恢复、复盘改进与报告五个阶段。'
            '事件发现来源包括监控告警、用户报障、巡检发现、安全告警等。定级后 5 分钟内向采购人相关联系人'
            '通报。技术处置阶段优先恢复业务，再分析根因。业务恢复后进行验证，确认功能、性能、数据一致性均'
            '正常。复盘阶段输出《事件复盘报告》，提出改进项并跟踪闭环。图 6 展示了应急响应流程。'
        )
        self._add_picture('应急响应流程图', 6, width=Inches(6.2))

        self.add_heading3('1.7.1 重要时期保障')
        self.add_body_paragraph(
            '针对重大会议、节假日、重要政务活动等关键时段，我方启动重要时期保障机制。保障机制包括：'
            '提前两周开展风险评估与隐患排查，提前一周完成配置冻结与备份，保障期间实行 7x24 加强值守，'
            '保障结束后开展复盘总结。保障期间原则上禁止非紧急变更，所有操作均需双人复核。'
        )

        # 1.8 培训与知识转移
        self.add_heading2('1.8 培训与知识转移方案')
        self.add_body_paragraph(
            '我方将培训与知识转移视为服务保障可持续的重要支撑。培训体系覆盖管理员、运维人员、最终用户'
            '三个群体，内容涵盖政务微信日常管理、智能网关运维、信息安全与保密、应急与业务连续性、'
            '数据运营与看板、持续改进与知识管理六大课程。每门课程均明确目标、课时、大纲、考核方式与'
            '交付物。'
        )
        self.add_body_paragraph(
            '培训方式采用线上直播、线下授课、实操演练、桌面推演、视频回放等多种形式相结合。管理员培训'
            '侧重实操配置，考核方式为现场搭建；运维培训侧重故障模拟处置，考核方式为应急响应演练；用户'
            '培训侧重常用功能与自助排障，考核方式为在线问卷。培训完成后，我方提供培训教材、操作手册、'
            '视频资料与考核记录，供采购人后续自主使用。图 7 展示了培训与知识转移体系。'
        )
        self._add_picture('培训与知识转移体系', 7, width=Inches(6.2))

        # 1.9 验收与交付
        self.add_heading2('1.9 验收与交付安排')
        self.add_body_paragraph(
            '本项目验收分为月度检查、季度评估、年度终验三个层级。月度检查由项目经理组织，核查当月服务'
            '工单、巡检报告、变更记录、安全事件等；季度评估由项目总监牵头，对 KPI 达成情况、用户满意度、'
            '知识库增量、改进项闭环等进行综合评估；年度终验由采购人组织，我方提交年度服务总结、'
            '知识资产清单、满意度调查报告、SLA 达成报告、安全合规报告等，配合完成现场核查与问询。'
        )
        self.add_body_paragraph(
            '验收标准以招标文件与合同约定为准，重点包括：服务期内无重大安全保密事件、关键系统可用性达标、'
            'SLA 响应恢复时限达标率符合要求、用户满意度达到约定分值、培训与知识转移完成、知识库与'
            'CMDB 完整移交、所有报告按时提交。若某项指标未达标，我方承诺在约定时间内完成整改，'
            '并承担相应责任。'
        )

        # 1.10 逐条响应（扩写版）
        self.add_heading2('1.10 招标文件服务要求逐条响应')
        self.add_body_paragraph(
            '以下对我方理解与响应招标文件中的关键服务要求逐一说明。每一项均按照《招标要求》原文提炼、'
            '《我方响应》具体措施、《实施要点》落地路径、《交付物》核验依据四个维度展开，确保评委可快速'
            '对照、逐项核验。'
        )
        responses = [
            (
                '政务微信订阅许可续约',
                '确保 4300 个席位在合同年度内持续有效。',
                '建立许可到期预警机制，提前 30 天启动续约流程；维护席位台账，记录每个席位的部门、岗位、'
                '状态与到期时间；续约窗口期采用灰度切换，避免集中到期导致服务中断；续约完成后输出'
                '《许可续约报告》供采购人归档。',
                '《许可续约报告》《席位台账》《到期预警记录》'
            ),
            (
                '信创环境适配服务',
                '针对鲲鹏、飞腾、统信、麒麟、信创浏览器等开展兼容性测试、问题修复、回归验证。',
                '建立信创适配知识库，记录每类终端、操作系统、浏览器版本的兼容性状态；与厂商保持技术对接，'
                '获取补丁与更新信息；对新增问题执行发现、评估、修复、验证、上线、跟踪六步闭环；每季度输出'
                '《信创适配报告》。',
                '《信创适配报告》《信创知识库》《回归测试记录》'
            ),
            (
                '一年期软件使用许可',
                '保障软件功能使用权，完成账号配置、权限初始化、使用培训。',
                '梳理账号权限矩阵，按岗位分配默认角色；配置账号生命周期管理策略，包括开通、变更、冻结、'
                '注销；组织管理员培训，确保采购人可自主完成日常账号管理。',
                '《账号权限矩阵》《管理员培训记录》《使用许可清单》'
            ),
            (
                '一年期新版本升级许可',
                '在合同期内提供新版本的平滑升级服务。',
                '升级前开展影响评估，识别功能变更、接口变化、信创影响；在测试环境完成功能回归与压力测试；'
                '升级方案经审批后实施，优先采用灰度发布；升级后验证业务连续性并输出《升级验收报告》。',
                '《升级影响评估》《测试报告》《升级验收报告》'
            ),
            (
                '一年期云服务能力许可',
                '保障云端消息推送、直播、会议等云服务的调用额度与可用性。',
                '监控云服务配额使用率，设置阈值告警；对直播、会议等高并发场景提前评估带宽与并发容量；'
                '出现云服务异常时第一时间切换备用通道或降级策略。',
                '《云服务监控月报》《容量评估报告》《异常处置记录》'
            ),
            (
                '售后响应服务',
                '7x24 受理，工单、电话、远程多渠道支持，按一级/二级/三级分级响应与升级。',
                '服务台统一受理，30 秒内生成工单并分派；一级事件 15 分钟响应、2 小时恢复；二线专家 30 分钟'
                '内介入；必要时引入原厂三线支持；所有事件 24 小时内闭环归档。',
                '《工单记录》《事件处置报告》《SLA 达成报告》'
            ),
            (
                '远程软件授权服务',
                '负责授权策略下发、权限变更、失效回收。',
                '所有授权操作通过工单申请与审批，变更内容、审批人、执行人、时间全部留痕；每周核查授权状态，'
                '回收离职或调岗人员权限；对异常授权进行审计。',
                '《授权变更记录》《授权审计报告》《回收清单》'
            ),
            (
                '远程技术支持',
                '通过远程桌面、日志分析、配置核查等方式快速定位与处置问题。',
                '建立远程支持安全规范，所有远程操作需审批、双人在线、全程录屏；问题处置完成后更新知识库，'
                '对同类问题形成标准处置动作。',
                '《远程支持记录》《知识库条目》《操作录屏存档》'
            ),
            (
                '组织和用户运营服务',
                '定期梳理组织架构、治理通讯录、清理 inactive 账号、开展活跃度运营。',
                '每月同步组织架构，处理部门调整、人员异动；每季度清理 inactive 账号并通知对应部门；'
                '通过使用培训、推广活动提升活跃度；输出《组织运营月报》。',
                '《组织运营月报》《通讯录治理记录》《账号清理清单》'
            ),
            (
                '数据运营统计服务',
                '按月输出活跃、登录、功能使用、故障等维度统计看板与洞察报告。',
                '定义核心指标口径，包括活跃用户数、日登录人次、功能使用分布、消息发送量、故障工单量等；'
                '通过可视化看板与月度报告呈现趋势、异常与改进建议，支持管理决策。',
                '《数据运营月报》《指标口径说明》《看板截图》'
            ),
            (
                '应用运营服务',
                '负责自建应用上架审核、流程配置、消息推送、运营活动支撑。',
                '制定应用上架标准，包括安全审查、权限审查、用户体验审查；配置应用消息推送策略，避免打扰；'
                '对重要运营活动提供现场或远程保障。',
                '《应用上架审核单》《推送策略配置》《活动保障记录》'
            ),
            (
                '运行保障服务',
                '提供重要时期值守、全量巡检、隐患清零、双链路热备。',
                '制定年度巡检计划，覆盖平台、网关、应用、安全、数据各维度；重要时期提前两周进入保障状态，'
                '冻结非紧急变更，执行 7x24 加强值守；巡检发现的隐患纳入工单闭环。',
                '《巡检报告》《隐患闭环清单》《重要时期保障总结》'
            ),
            (
                '专家咨询服务',
                '每半年至少组织一次架构、安全、信创方向专家评审。',
                '根据项目阶段与突出问题确定评审主题；邀请内部资深专家与外部顾问参与；评审输出优化建议与'
                '改进计划，纳入下阶段服务计划跟踪。',
                '《专家评审报告》《改进计划跟踪表》'
            ),
            (
                '智能网关维保服务',
                '提供热线、现场、健康检查、远程支持、应急响应、配置资产管理。',
                '将网关纳入 CMDB，记录配置、版本、证书、关联关系；每月执行健康检查，包括性能、容量、安全、'
                '日志；发生故障时按 SLA 响应并恢复；应急响应后开展复盘。',
                '《网关健康检查报告》《CMDB 记录》《事件复盘报告》'
            ),
        ]
        for i, (req, short, impl, deliverables) in enumerate(responses, 1):
            self.add_heading3('1.10.%d %s' % (i, req))
            self.add_body_paragraph('【招标要求】%s' % req)
            self.add_body_paragraph('【我方响应】%s' % short)
            self.add_body_paragraph('【实施要点】%s' % impl)
            self.add_body_paragraph('【交付物】%s' % deliverables)
            self.add_body_paragraph(
                '【承诺】我方将以制度化、流程化、可度量的方式确保上述要求落地，并通过月度报告、'
                '季度评估、年度终验接受采购人核验。'
            )

        # 1.11 落地保障详述
        self.add_heading2('1.11 落地保障详述与实施要点')
        self.add_body_paragraph(
            '为确保前述技术方案、管理制度、SLA 承诺真正落地，我方在组织、流程、工具、人员、沟通五个方面'
            '建立配套保障机制。以下从实际操作角度详细阐述每项保障机制的具体做法、责任人与交付频率，'
            '使方案不仅停留在纸面，更能转化为可执行、可检查、可改进的服务行动。'
        )

        self.add_heading3('1.11.1 组织保障')
        self.add_body_paragraph(
            '项目组织保障的核心是清晰的角色定义与稳定的团队配置。项目总监由公司高级管理者担任，'
            '对项目整体目标、客户满意度、重大风险负最终责任；项目经理全职投入，负责日常计划、资源协调、'
            '进度跟踪与报告输出；技术负责人负责平台、网关、应用的技术方向与疑难问题决策；安全负责人'
            '负责等保、保密、合规与审计事务；运营负责人负责用户运营、应用运营与数据运营。'
        )
        self.add_body_paragraph(
            '关键岗位设置 AB 角。项目经理的 B 角由资深服务主管担任，可在项目经理请假或突发缺位时立即'
            '接管；技术负责人设置网关、平台、应用三个技术方向的专项负责人，互为 backup；应急响应岗位'
            '实行轮值制，确保任何时刻均有熟悉项目的人员在岗。团队稳定性通过长期服务合同、绩效考核、'
            '知识沉淀三重机制保证，人员变动率控制在年度 10% 以内。'
        )
        self.add_body_paragraph(
            '沟通机制上，我方建立三级例会：日站会同步当日工单与风险；周例会回顾本周 KPI、巡检、变更、'
            '事件；月度例会汇报 SLA 达成、满意度、改进项、下月计划。所有会议均有纪要，并通过邮件与项目'
            '微信群同步采购人。重大事项 15 分钟内电话或微信通报，确保信息不延误。'
        )

        self.add_heading3('1.11.2 流程保障')
        self.add_body_paragraph(
            '流程保障的目标是让所有服务动作有章可循、有记录可查。我方为项目建立七大核心流程：服务台与'
            '事件管理流程、问题管理流程、变更管理流程、配置管理流程、发布管理流程、知识管理流程、'
            '服务请求管理流程。每个流程均包含流程图、角色职责、输入输出、关键控制点、SLA 时限、'
            '常见场景模板。'
        )
        self.add_body_paragraph(
            '事件管理流程示例：用户通过热线、工单、微信报障后，服务台 5 分钟内完成受理并记录；'
            '根据影响范围与紧急程度确定事件等级；一级事件 15 分钟内响应、2 小时内恢复；处置过程中每 30 分钟'
            '向采购人通报进展；恢复后进行验证并关闭工单；24 小时内输出事件报告。问题管理流程则针对'
            '重复或重大事件开展根因分析，输出预防措施并跟踪实施。'
        )
        self.add_body_paragraph(
            '变更管理流程对生产环境的所有配置、版本、架构调整进行管控。变更分为标准变更、常规变更、'
            '紧急变更三类。标准变更按预授权清单执行；常规变更需提交变更申请，经评估、审批、测试后实施；'
            '紧急变更可事后补办审批，但须保留完整处置记录。所有变更均有回滚方案，变更窗口避开业务高峰。'
        )

        self.add_heading3('1.11.3 工具保障')
        self.add_body_paragraph(
            '工具保障提升服务效率与可追溯性。我方部署统一运维管理平台，集成工单系统、监控系统、'
            '知识库系统、CMDB、报表系统。工单系统记录所有服务请求、事件、问题、变更，支持按时间、'
            '类型、状态、满意度多维度检索；监控系统对平台、网关、链路、云服务进行 7x24 监控，'
            '设置多级阈值告警；知识库系统沉淀 SOP、故障案例、适配经验、培训教材；CMDB 记录软硬件资产、'
            '配置项、关联关系；报表系统自动生成月度、季度、年度报告。'
        )
        self.add_body_paragraph(
            '监控指标体系覆盖可用性、性能、容量、安全四个维度。可用性指标包括系统整体可用率、核心接口'
            '成功率、网关连通率；性能指标包括平均响应时间、并发处理能力、消息推送延迟；容量指标包括'
            '席位使用率、云资源配额、存储空间；安全指标包括未授权访问尝试、漏洞修复率、账号异常行为。'
            '所有指标均有阈值、告警与报告机制。'
        )
        self.add_body_paragraph(
            '数据安全方面，工具平台采用角色分级访问，敏感操作需二次认证；所有日志留存不少于六个月；'
            '知识库中的敏感案例需脱敏处理；CMDB 与工单数据定期备份。我方承诺不向第三方泄露任何项目数据，'
            '离场时完整移交数据与权限。'
        )

        self.add_heading3('1.11.4 人员与能力保障')
        self.add_body_paragraph(
            '人员与能力保障是服务质量的基础。项目经理具备 PMP 或同等项目管理资质，并拥有五年以上政务'
            '协同项目经验；技术负责人具备平台运维、网关运维、安全合规相关认证；驻场工程师具备一年以上的'
            '同类项目经验并通过本项目专项培训。所有人员上岗前签署《保密协议》与《服务承诺书》。'
        )
        self.add_body_paragraph(
            '培训体系分为入职培训、专项培训、持续培训三类。入职培训包括公司制度、项目背景、安全保密；'
            '专项培训针对政务微信、智能网关、信创环境、应急响应开展；持续培训通过月度技术分享、案例复盘、'
            '厂商交流保持团队能力更新。培训记录纳入人员档案，作为绩效考核依据。'
        )
        self.add_body_paragraph(
            '能力提升还与知识管理联动。每季度评选优秀案例、最佳 SOP、最快恢复奖，激励团队沉淀经验；'
            '对复杂故障开展复盘会，形成根因分析与改进项；对新型技术问题引入厂商专家进行专题培训。'
            '通过培训、复盘、激励三位一体，确保团队能力持续匹配项目需求。'
        )

        self.add_heading3('1.11.5 沟通与报告保障')
        self.add_body_paragraph(
            '沟通与报告保障让采购人始终掌握服务状态。日常沟通渠道包括：项目经理专线电话、微信服务群、'
            '工单系统、月度例会。紧急事件通过电话或微信即时通报；一般问题通过工单系统跟踪；定期报告通过'
            '月度报告、季度评估、年度总结呈现。'
        )
        self.add_body_paragraph(
            '月度报告内容包括：服务概览、工单统计、事件分析、巡检结果、变更记录、KPI 达成、改进项进展、'
            '下月计划。季度评估报告增加用户满意度分析、知识库增量、培训情况、专家评审结论。年度总结报告'
            '包括全年服务回顾、SLA 总达成、重大事件复盘、知识资产清单、续约建议。所有报告均提供电子版与'
            '盖章纸质版。'
        )
        self.add_body_paragraph(
            '此外，我方建立服务满意度调查机制。每季度向采购人代表与关键用户发放满意度问卷，调查内容包括'
            '响应速度、解决效果、沟通态度、报告质量、培训效果等。满意度低于 90 分时，项目经理需在 5 个'
            '工作日内提交改进计划，并跟踪至达标。'
        )

        # 1.12 案例应用与经验沉淀
        self.add_heading2('1.12 典型案例应用与经验沉淀')
        self.add_body_paragraph(
            '以下典型案例进一步说明我方如何将前述方法论、流程、工具应用于实际服务场景。每个案例均来自'
            '同类政务协同项目经验，已做脱敏处理，重点展示问题诊断、处置过程、经验沉淀与持续改进。'
        )
        cases = [
            ('案例一：许可集中到期导致登录受限', '某单位因许可到期前未收到预警，导致部分账号登录受限。'
             '我方接手后，建立许可到期预警机制，提前 30 天通过短信、邮件、工单通知采购人；梳理席位台账，'
             '按部门与到期时间排序；续约窗口采用分批灰度切换，避免集中操作风险。后续未再发生类似事件。'),
            ('案例二：信创终端应用白屏', '某街道使用信创终端打开政务微信应用出现白屏。'
             '我方通过远程诊断确认该终端 WebView 版本与前端框架不兼容，协调终端厂商推送补丁，并在测试环境'
             '验证后灰度发布。该案例纳入信创适配知识库，后续同类问题处置时间从 3 小时缩短至 30 分钟。'),
            ('案例三：网关链路路由震荡', '智能网关外网访问偶发中断，监控显示丢包率升高。'
             '我方第一时间切换备用链路，并联合运营商定位路由震荡节点，推动优化。后续增加链路质量监控与'
             '自动切换策略，中断恢复时间从 2 小时降至 5 分钟。'),
            ('案例四：重要会议直播保障', '大型政务会议期间，直播并发突增，部分分会场卡顿。'
             '我方提前扩容云端带宽，限制非关键流量，并安排专人值守。会议全程无重大中断，结束后输出保障'
             '复盘，优化容量评估模型。'),
            ('案例五：权限误删引发批量影响', '某角色被误删导致多个部门功能受限。'
             '我方从权限备份中恢复，并补办变更审批流程。后续引入权限变更双人复核与自动化备份机制，'
             '类似事件零复发。'),
            ('案例六：数据看板延迟影响决策', '月度统计看板数据延迟一天，影响管理决策。'
             '我方定位到 ETL 任务资源争用，调整错峰调度并增加告警。后续看板数据在每日凌晨 6 点前完成更新，'
             '并通过监控日报确认。'),
        ]
        for title, content in cases:
            self.add_heading3(title)
            self.add_body_paragraph(content)
            self.add_body_paragraph(
                '经验沉淀：该案例已转化为知识库条目、SOP 动作或检查单项目，纳入定期培训与月度复盘，'
                '确保同类问题可快速复用处置经验。'
            )

    def generate_full_template(self, output_path):
        """完整响应文件：前件 + 技术方案概述 + 详细技术方案（Fb 母版）+ 附录。"""
        self.doc = Document()
        self.setup_page()
        self.setup_styles()
        # 响应文件前件（封面、投标函、各类一览表、授权书、身份证明）
        self.generate_front_matter()
        # 目录占位
        self.generate_toc()
        # 技术方案概述（带图片）
        self.generate_technical_proposal()
        # 详细技术方案与全部附录（由 FbBidGenerator 继承链完整生成）
        super().generate_section1_proposal()
        self.doc.save(output_path)


def main():
    import argparse
    parser = argparse.ArgumentParser(description="生成盐田政务微信项目响应文件格式完整技术标")
    parser.add_argument("--output", default="盐田区政务微信授权智能网关服务保障项目投标书(响应文件格式版).docx")
    args = parser.parse_args()
    generator = ResponseBidGenerator()
    generator.generate_full_template(args.output)
    print("已生成（响应文件格式版）：", args.output)


if __name__ == '__main__':
    main()
